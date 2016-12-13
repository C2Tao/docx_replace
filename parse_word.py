from docx import Document
from openpyxl import load_workbook
import os, re

def replace_pattern(document, old_text, new_text):
    for paragraph in document.paragraphs:
        for inline in paragraph.runs:
            inline.text = re.sub(old_text, new_text, inline.text)

    for table in document.tables:
        for r in range(len(table.rows)):
            for cell in table.row_cells(r):
                for paragraph in cell.paragraphs:
                    for inline in paragraph.runs:
                        inline.text = re.sub(old_text, new_text, inline.text)


def read_excel(dat_file):
    wb = load_workbook(filename = dat_file)
    ws = wb.get_active_sheet()
    data = {}
    for i in range(ws.max_column):
        a = chr(ord('A') + i)
        if not ws[a+'1'].value: break
        for j in range(ws.max_row):
            ai = a+str(j+1)
            cell = ws[ai].value
            if not cell: break
            if j==0:
                attr = cell
                data[attr] = []
            else:
                try: 
                    data[attr].append(str(cell))
                except:
                    data[attr].append(cell)
    return data

def write_word(old_file, dat_file, new_file):
    data = read_excel(dat_file)
    for i in range(len(data.values()[0])):
        document = Document(old_file)
        for k in data.keys():
            replace_pattern(document, k, data[k][i])
        document.save(new_file.format(i))

#path = r'C:\\Users\\c2tao\\Desktop\\word\\'
#path ='/mnt/c/Users/c2tao/Desktop/word/'
path = 'word/'
if not os.path.exists(path): os.makedirs(path)

new_file = path+'new-file-name-{}.docx'
dat_file = 'database.xlsx'
old_file = 'template.docx'
write_word(old_file, dat_file, new_file)
